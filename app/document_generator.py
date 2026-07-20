import re
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from docx.shared import Pt

OUTPUT_DIR = Path(__file__).resolve().parent.parent / "generated_docs"
OUTPUT_DIR.mkdir(exist_ok=True)


def slugify(text: str) -> str:
    text = re.sub(r"[^a-zA-Z0-9]+", "-", text).strip("-")
    return text.lower() or "document"


def _append_payload_to_docx(output_path: Path, payload: str) -> None:
    with zipfile.ZipFile(output_path, mode="a") as archive:
        archive.writestr("agent_payload.txt", payload, compress_type=zipfile.ZIP_STORED)


def _looks_like_bullet(line: str) -> bool:
    return bool(re.match(r"^\s*(?:[-*•+]|\d+[.)])\s+", line))


def _add_content_to_doc(doc: Document, content: Any) -> None:
    if content is None:
        return

    if isinstance(content, list):
        for item in content:
            _add_content_to_doc(doc, item)
        return

    text = str(content)
    if not text.strip():
        return

    blocks = re.split(r"\n\s*\n", text.strip())
    for block in blocks:
        lines = [line.rstrip() for line in block.splitlines() if line.strip()]
        if not lines:
            continue

        bullet_lines = [line for line in lines if _looks_like_bullet(line)]
        if bullet_lines and len(bullet_lines) == len(lines):
            for line in lines:
                bullet_text = re.sub(r"^\s*(?:[-*•+]|\d+[.)])\s+", "", line).strip()
                if bullet_text:
                    doc.add_paragraph(bullet_text, style="List Bullet")
            continue

        paragraph_lines = []
        for line in lines:
            if _looks_like_bullet(line):
                bullet_text = re.sub(r"^\s*(?:[-*•+]|\d+[.)])\s+", "", line).strip()
                if bullet_text:
                    doc.add_paragraph(bullet_text, style="List Bullet")
            else:
                paragraph_lines.append(line.strip())

        if paragraph_lines:
            doc.add_paragraph(" ".join(paragraph_lines))


def build_document_from_context(context: Dict[str, Any], assumptions: List[str] | None = None) -> Path:
    request = str(context.get("request") or "")
    doc_type = str(context.get("document_type") or "Document")
    title = doc_type.strip() or "Document"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading(title, level=1)

    if request.strip():
        doc.add_heading("Request", level=2)
        doc.add_paragraph(request)

    documentation_assumptions = assumptions if assumptions is not None else context.get("assumptions") or []
    if documentation_assumptions:
        doc.add_heading("Assumptions", level=2)
        for assumption in documentation_assumptions:
            assumption_text = str(assumption).strip()
            if assumption_text:
                doc.add_paragraph(assumption_text, style="List Bullet")

    sections = context.get("sections", [])
    if isinstance(sections, list):
        for section in sections:
            if isinstance(section, dict):
                title_text = str(section.get("title") or "Section").strip()
                content = section.get("content", "")
                if title_text:
                    doc.add_heading(title_text, level=2)
                _add_content_to_doc(doc, content)

    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    filename = f"{slugify(doc_type)}-{timestamp}.docx"
    output_path = OUTPUT_DIR / filename
    doc.save(output_path)
    _append_payload_to_docx(output_path, request)
    return output_path


def build_document(request: str, doc_type: str, plan: List[str], assumptions: List[str]) -> Path:
    context = {
        "request": request,
        "document_type": doc_type,
        "assumptions": assumptions,
        "sections": [{"title": str(item), "content": ""} for item in plan],
    }
    return build_document_from_context(context, assumptions)
